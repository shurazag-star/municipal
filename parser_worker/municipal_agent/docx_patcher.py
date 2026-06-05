from __future__ import annotations

from collections import defaultdict
from copy import deepcopy
from decimal import Decimal, ROUND_HALF_UP
from pathlib import Path
from typing import Any, Dict, Iterable, List
import re

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import _Row


def format_money_for_docx(
    amount_rub: Decimal | str | int | float,
    source_cell_raw_value: str,
    unit: str = "thousand_rub",
    default_grouping: bool = False,
) -> str:
    amount = Decimal(str(amount_rub))
    if unit == "thousand_rub":
        amount = amount / Decimal("1000")

    decimal_places = _decimal_places(source_cell_raw_value)
    if amount and amount.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP) != amount.quantize(Decimal("0.1"), rounding=ROUND_HALF_UP):
        decimal_places = max(decimal_places, 2)
    quant = Decimal("1") if decimal_places == 0 else Decimal("1").scaleb(-decimal_places)
    amount = amount.quantize(quant, rounding=ROUND_HALF_UP)
    grouping_separator = _grouping_separator(source_cell_raw_value)
    if grouping_separator is None and default_grouping and abs(amount) >= Decimal("1000"):
        grouping_separator = " "
    if grouping_separator:
        formatted = f"{amount:,.{decimal_places}f}"
        return formatted.replace(",", grouping_separator).replace(".", ",")

    return f"{amount:.{decimal_places}f}".replace(".", ",")


def patch_docx(
    input_path: str | Path,
    output_path: str | Path,
    changes: Iterable[Dict[str, Any]] | Dict[str, Any],
) -> Dict[str, Any]:
    document = Document(str(input_path))
    payload = _normalize_patch_payload(changes)
    applied: List[Dict[str, Any]] = []
    skipped: List[Dict[str, Any]] = []

    for change in payload["cell_updates"]:
        try:
            table_index = int(change["table_index"])
            row_index = int(change["row_index"])
            cell_index = int(change["cell_index"])
            cell = _visual_cell(document, table_index, row_index, cell_index)
            raw_value = str(change.get("source_cell_raw_value") or cell.text)
            formatted_value = format_money_for_docx(
                change["amount_rub"],
                source_cell_raw_value=raw_value,
                unit=change.get("unit") or "thousand_rub",
            )
            _set_cell_text_preserving_first_run(cell, formatted_value)
            applied.append(
                {
                    "change_item_id": change.get("change_item_id"),
                    "table_index": table_index,
                    "row_index": row_index,
                    "cell_index": cell_index,
                    "value": formatted_value,
                }
            )
        except (IndexError, KeyError, TypeError, ValueError) as error:
            skipped.append(
                {
                    "change_item_id": change.get("change_item_id"),
                    "table_index": change.get("table_index"),
                    "row_index": change.get("row_index"),
                    "cell_index": change.get("cell_index"),
                    "error": str(error),
                }
            )

    text_result = _apply_text_updates(document, payload["text_updates"])
    result_count_result = _update_result_counts_for_insertions(document, payload["insert_objects"])
    insertion_result = _insert_objects(document, payload["insert_objects"])
    approval_result = _normalize_approval_header(document)
    document.save(str(output_path))
    return {
        "output_path": str(output_path),
        "applied_count": len(applied),
        "skipped_count": len(skipped),
        "text_applied_count": len(text_result["applied"]),
        "text_skipped_count": len(text_result["skipped"]),
        "result_count_applied_count": len(result_count_result["applied"]),
        "result_count_skipped_count": len(result_count_result["skipped"]),
        "applied": applied,
        "skipped": skipped,
        "text_applied": text_result["applied"],
        "text_skipped": text_result["skipped"],
        "result_count_applied": result_count_result["applied"],
        "result_count_skipped": result_count_result["skipped"],
        "inserted_count": len(insertion_result["inserted"]),
        "inserted_rows_count": insertion_result["inserted_rows_count"],
        "skipped_insertions_count": len(insertion_result["skipped_insertions"]),
        "inserted": insertion_result["inserted"],
        "skipped_insertions": insertion_result["skipped_insertions"],
        "approval_header_normalized": approval_result["normalized"],
    }


def _normalize_patch_payload(changes: Iterable[Dict[str, Any]] | Dict[str, Any]) -> Dict[str, List[Dict[str, Any]]]:
    if isinstance(changes, dict):
        return {
            "cell_updates": list(changes.get("cell_updates") or []),
            "text_updates": list(changes.get("text_updates") or []),
            "insert_objects": list(changes.get("insert_objects") or []),
        }

    return {
        "cell_updates": list(changes or []),
        "text_updates": [],
        "insert_objects": [],
    }


def _visual_cell(document, table_index: int, row_index: int, cell_index: int):
    # Use the same visual grid that docx_parser reads via row.cells. python-docx
    # table.cell(row, col) can resolve to a wrong physical cell when a table mixes
    # horizontal gridSpan cells with vertical merges, which is common in municipal
    # finance tables.
    return document.tables[table_index].rows[row_index].cells[cell_index]


def _apply_text_updates(document, text_updates: List[Dict[str, Any]]) -> Dict[str, List[Dict[str, Any]]]:
    applied: List[Dict[str, Any]] = []
    skipped: List[Dict[str, Any]] = []
    for change in text_updates:
        try:
            table_index = int(change["table_index"])
            row_index = int(change["row_index"])
            cell_index = int(change["cell_index"])
            cell = _visual_cell(document, table_index, row_index, cell_index)
            value = str(change.get("text") or "")
            _set_cell_text_preserving_first_run(cell, value)
            applied.append(
                {
                    "table_index": table_index,
                    "row_index": row_index,
                    "cell_index": cell_index,
                    "value": value,
                }
            )
        except (IndexError, KeyError, TypeError, ValueError) as error:
            skipped.append(
                {
                    "table_index": change.get("table_index"),
                    "row_index": change.get("row_index"),
                    "cell_index": change.get("cell_index"),
                    "error": str(error),
                }
            )

    return {"applied": applied, "skipped": skipped}


def _insert_objects(document, insert_objects: List[Dict[str, Any]]) -> Dict[str, Any]:
    inserted: List[Dict[str, Any]] = []
    skipped: List[Dict[str, Any]] = []
    inserted_rows_count = 0
    grouped: dict[tuple[int, int], list[tuple[int, Dict[str, Any]]]] = defaultdict(list)

    for sequence, spec in enumerate(insert_objects):
        try:
            table_index = int(spec["table_index"])
            insert_after_row_index = int(spec["insert_after_row_index"])
            grouped[(table_index, insert_after_row_index)].append((sequence, spec))
        except (KeyError, TypeError, ValueError) as error:
            skipped.append(
                {
                    "change_item_ids": spec.get("change_item_ids"),
                    "error": str(error),
                }
            )

    for (table_index, insert_after_row_index), group in sorted(grouped.items(), key=lambda item: item[0], reverse=True):
        try:
            table = document.tables[table_index]
            anchor_row = table.rows[insert_after_row_index]
        except (IndexError, TypeError) as error:
            for _sequence, spec in group:
                skipped.append(
                    {
                        "change_item_ids": spec.get("change_item_ids"),
                        "table_index": spec.get("table_index"),
                        "insert_after_row_index": spec.get("insert_after_row_index"),
                        "error": str(error),
                    }
                )
            continue

        anchor_tr = anchor_row._tr
        for _sequence, spec in sorted(group, key=lambda item: item[0]):
            try:
                spec = dict(spec)
                if not _clean_text(spec.get("responsible")):
                    inferred_responsible = _neighbor_responsible_for_insert(table, insert_after_row_index, spec)
                    if inferred_responsible:
                        spec["responsible"] = inferred_responsible
                template_row_index = int(spec.get("template_row_index") or insert_after_row_index)
                template_tr = deepcopy(table.rows[template_row_index]._tr)
                object_rows = list(spec.get("rows") or [])
                if not object_rows:
                    object_rows = [{"source_type": "TOTAL", "source_label": "Итого", "amounts_by_year": {}}]

                object_inserted_rows = []
                inserted_row_objects = []
                for row_spec in object_rows:
                    new_tr = deepcopy(template_tr)
                    _detach_vertical_merges(new_tr)
                    anchor_tr.addnext(new_tr)
                    anchor_tr = new_tr
                    row = _Row(new_tr, table)
                    _populate_inserted_row(row, spec, row_spec)
                    inserted_row_objects.append(row)
                    object_inserted_rows.append(
                        {
                            "source_type": row_spec.get("source_type"),
                            "source_label": row_spec.get("source_label"),
                        }
                    )
                    inserted_rows_count += 1

                merged_identity_cells = _merge_inserted_object_identity_cells(inserted_row_objects)
                inserted.append(
                    {
                        "change_item_ids": spec.get("change_item_ids") or [],
                        "table_index": table_index,
                        "insert_after_row_index": insert_after_row_index,
                        "rows_count": len(object_inserted_rows),
                        "merged_identity_cells": merged_identity_cells,
                        "object_name": spec.get("object_name"),
                    }
                )
            except (IndexError, KeyError, TypeError, ValueError) as error:
                skipped.append(
                    {
                        "change_item_ids": spec.get("change_item_ids"),
                        "table_index": spec.get("table_index"),
                        "insert_after_row_index": spec.get("insert_after_row_index"),
                        "template_row_index": spec.get("template_row_index"),
                        "error": str(error),
                    }
                )

    return {
        "inserted": inserted,
        "inserted_rows_count": inserted_rows_count,
        "skipped_insertions": skipped,
    }


def _update_result_counts_for_insertions(document, insert_objects: List[Dict[str, Any]]) -> Dict[str, List[Dict[str, Any]]]:
    applied: List[Dict[str, Any]] = []
    skipped: List[Dict[str, Any]] = []
    for spec in insert_objects:
        try:
            table_index = int(spec["table_index"])
            table = document.tables[table_index]
            parent_display = _normalize_display_number(str(spec.get("parent_display_number") or ""))
            completion_year = _completion_year(spec)
            if not parent_display or completion_year is None:
                continue

            row_index = _find_result_count_row(table, parent_display)
            if row_index is None:
                continue

            row = table.rows[row_index]
            year_by_col = _result_year_columns(table, row_index)
            changed_cells = []
            original_values = {
                col: _integer_cell_value(row.cells[col])
                for col in range(len(row.cells))
            }
            if _set_incremented_integer_cell(row.cells[4], original_values.get(4)):
                changed_cells.append(4)
            for col, year in year_by_col.items():
                original_value = original_values.get(col)
                if year == completion_year and col < len(row.cells) and original_value and _set_incremented_integer_cell(row.cells[col], original_value):
                    changed_cells.append(col)
            if changed_cells:
                applied.append(
                    {
                        "table_index": table_index,
                        "row_index": row_index,
                        "parent_display_number": spec.get("parent_display_number"),
                        "completion_year": completion_year,
                        "cell_indices": changed_cells,
                    }
                )
        except (IndexError, KeyError, TypeError, ValueError) as error:
            skipped.append(
                {
                    "table_index": spec.get("table_index"),
                    "parent_display_number": spec.get("parent_display_number"),
                    "error": str(error),
                }
            )

    return {"applied": applied, "skipped": skipped}


def _completion_year(spec: Dict[str, Any]) -> int | None:
    years = []
    for value in spec.get("active_years") or []:
        try:
            years.append(int(value))
        except (TypeError, ValueError):
            continue
    if years:
        return max(years)

    match = re.findall(r"20\d{2}", str(spec.get("execution_period") or ""))
    return int(match[-1]) if match else None


def _find_result_count_row(table, parent_display: str) -> int | None:
    candidates = {_normalize_display_number(parent_display), _normalize_display_number(f"{parent_display}.")}
    for row_index, row in enumerate(table.rows):
        cells = row.cells
        if len(cells) < 5:
            continue
        display = _normalize_display_number(cells[0].text)
        if display not in candidates:
            continue
        if _clean_text(cells[2].text).upper() != "Х" or _clean_text(cells[3].text).upper() != "Х":
            continue
        if _integer_cell_value(cells[4]) is None:
            continue
        return row_index
    return None


def _result_year_columns(table, row_index: int) -> Dict[int, int]:
    if row_index <= 0:
        return {}
    header_rows = [table.rows[index] for index in range(max(0, row_index - 2), row_index)]
    result: Dict[int, int] = {}
    current_year = None
    max_cells = max((len(row.cells) for row in header_rows), default=0)
    for col in range(5, max_cells):
        explicit_year = None
        for header in reversed(header_rows):
            if col >= len(header.cells):
                continue
            match = re.search(r"20\d{2}", _clean_text(header.cells[col].text))
            if match:
                explicit_year = int(match.group(0))
                break
        if explicit_year is not None:
            current_year = explicit_year
        if current_year is not None:
            result[col] = current_year
    return result


def _integer_cell_value(cell) -> int | None:
    text = _clean_text(cell.text)
    if not re.fullmatch(r"-?\d+", text):
        return None
    return int(text)


def _increment_integer_cell(cell) -> bool:
    value = _integer_cell_value(cell)
    if value is None:
        return False
    _set_cell_text_preserving_first_run(cell, str(value + 1))
    return True


def _set_incremented_integer_cell(cell, original_value: int | None) -> bool:
    if original_value is None:
        return False
    _set_cell_text_preserving_first_run(cell, str(original_value + 1))
    return True


def _normalize_display_number(value: str) -> str:
    return _clean_text(value).rstrip(".")


def _clean_text(value: str) -> str:
    return " ".join(str(value or "").replace("\u00a0", " ").replace("\u202f", " ").split())


def _populate_inserted_row(row, object_spec: Dict[str, Any], row_spec: Dict[str, Any]) -> None:
    cells = row.cells
    raw_cell_values = {index: cell.text for index, cell in enumerate(cells)}
    display_number = _display_number_for_inserted_row(object_spec, raw_cell_values.get(0, ""))
    _set_cell_if_present(cells, 0, display_number)
    _set_cell_if_present(cells, 1, str(object_spec.get("object_name") or ""))
    _set_cell_if_present(cells, 2, str(object_spec.get("execution_period") or ""))
    _set_cell_if_present(cells, 3, str(row_spec.get("source_label") or row_spec.get("source_type") or ""))

    total_cell_index = int(object_spec.get("total_cell_index") or 4)
    for index in range(total_cell_index, len(cells)):
        _set_cell_text_preserving_first_run(cells[index], "")

    unit = row_spec.get("unit") or object_spec.get("unit") or "thousand_rub"
    if "total_amount_rub" in row_spec:
        _set_money_cell(cells, total_cell_index, row_spec["total_amount_rub"], unit, raw_cell_values.get(total_cell_index, ""))

    year_cell_indices = object_spec.get("year_cell_indices") or {}
    amounts_by_year = row_spec.get("amounts_by_year") or {}
    for year, cell_index in year_cell_indices.items():
        amount = amounts_by_year.get(str(year), amounts_by_year.get(int(year), "0"))
        _set_money_cell(cells, int(cell_index), amount, unit, raw_cell_values.get(int(cell_index), ""))

    responsible = _clean_text(object_spec.get("responsible"))
    responsible_cell_index = _responsible_cell_index(cells, object_spec)
    if responsible and responsible_cell_index is not None:
        _set_cell_text_preserving_first_run(cells[responsible_cell_index], responsible)


def _neighbor_responsible_for_insert(table, insert_after_row_index: int, object_spec: Dict[str, Any]) -> str:
    column_index = _responsible_cell_index(table.rows[insert_after_row_index].cells, object_spec)
    if column_index is None:
        return ""

    for row_index in range(insert_after_row_index + 1, len(table.rows)):
        responsible = _responsible_text_from_row(table.rows[row_index], column_index)
        if responsible:
            return responsible

    for row_index in range(insert_after_row_index, -1, -1):
        responsible = _responsible_text_from_row(table.rows[row_index], column_index)
        if responsible:
            return responsible

    return ""


def _responsible_text_from_row(row, column_index: int) -> str:
    if column_index >= len(row.cells):
        return ""

    text = _clean_text(row.cells[column_index].text)
    return "" if text.upper() == "Х" else text


def _responsible_cell_index(cells, object_spec: Dict[str, Any]) -> int | None:
    explicit = object_spec.get("responsible_cell_index")
    if explicit is not None:
        try:
            index = int(explicit)
            return index if 0 <= index < len(cells) else None
        except (TypeError, ValueError):
            return None

    try:
        total_cell_index = int(object_spec.get("total_cell_index") or 4)
    except (TypeError, ValueError):
        total_cell_index = 4
    amount_indexes = [total_cell_index]
    for value in (object_spec.get("year_cell_indices") or {}).values():
        try:
            amount_indexes.append(int(value))
        except (TypeError, ValueError):
            continue

    candidate = len(cells) - 1
    return candidate if candidate > max(amount_indexes, default=total_cell_index) else None


def _normalize_approval_header(document) -> Dict[str, bool]:
    if not document.tables:
        return {"normalized": False}

    try:
        cell = document.tables[0].rows[0].cells[0]
    except IndexError:
        return {"normalized": False}

    text = cell.text
    normalized = re.sub(
        r"от\s+\d{1,2}\.\d{1,2}\.\d{4}\s*(?:г\.?)?\s*№\s*[\wА-Яа-яЁё./\\-]+",
        "от _______________ №__________",
        text,
        count=1,
    )
    if normalized == text:
        return {"normalized": False}

    _set_cell_text_preserving_first_run(cell, normalized)
    return {"normalized": True}


def _merge_inserted_object_identity_cells(rows: List[_Row]) -> List[int]:
    if len(rows) < 2:
        return []

    merged_columns: List[int] = []
    for column_index in (0, 1, 2):
        try:
            for row in rows[1:]:
                _set_cell_if_present(row.cells, column_index, "")
            _set_vertical_merge(rows[0].cells[column_index], "restart")
            for row in rows[1:]:
                _set_vertical_merge(row.cells[column_index], "continue")
            merged_columns.append(column_index)
        except (IndexError, ValueError):
            continue
    return merged_columns


def _set_vertical_merge(cell, value: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    vertical_merge = properties.find(qn("w:vMerge"))
    if vertical_merge is None:
        vertical_merge = OxmlElement("w:vMerge")
        properties.append(vertical_merge)

    value_key = qn("w:val")
    if value == "restart":
        vertical_merge.set(value_key, "restart")
    elif value_key in vertical_merge.attrib:
        del vertical_merge.attrib[value_key]


def _display_number_for_inserted_row(object_spec: Dict[str, Any], template_value: str) -> str:
    value = str(object_spec.get("display_number") or "")
    if str(template_value or "").strip().endswith(".") and value and not value.endswith("."):
        return f"{value}."
    return value


def _detach_vertical_merges(table_row) -> None:
    # A cloned row can carry w:vMerge continuation/restart markers from the
    # source object's multi-row block. If the inserted row keeps them, Word and
    # python-docx can bind the new row to the previous object and the re-parser
    # will attribute old funding rows to the new object.
    for cell in table_row.tc_lst:
        properties = cell.tcPr
        if properties is None:
            continue
        vertical_merge = properties.find(qn("w:vMerge"))
        if vertical_merge is not None:
            properties.remove(vertical_merge)


def _set_cell_if_present(cells, index: int, value: str) -> None:
    if 0 <= index < len(cells):
        _set_cell_text_preserving_first_run(cells[index], value)


def _set_money_cell(cells, index: int, amount: Decimal | str | int | float, unit: str, raw_value: str) -> None:
    if 0 <= index < len(cells):
        formatted = format_money_for_docx(amount, source_cell_raw_value=raw_value, unit=unit, default_grouping=True)
        _set_cell_text_preserving_first_run(cells[index], formatted)


def _decimal_places(raw_value: str) -> int:
    cleaned = str(raw_value or "").strip()
    match = re.search(r"[,\.](\d+)\s*$", cleaned)
    return len(match.group(1)) if match else 0


def _grouping_separator(raw_value: str) -> str | None:
    text = str(raw_value or "")
    for separator in ("\u00a0", "\u202f", " "):
        if re.search(rf"\d{re.escape(separator)}\d{{3}}", text):
            return separator
    return None


def _set_cell_text_preserving_first_run(cell, value: str) -> None:
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    if paragraph.runs:
        paragraph.runs[0].text = value
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(value)

    for extra_paragraph in cell.paragraphs[1:]:
        for run in extra_paragraph.runs:
            run.text = ""
