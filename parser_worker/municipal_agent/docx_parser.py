from __future__ import annotations

from dataclasses import dataclass, field
from decimal import Decimal
from pathlib import Path
import re
from typing import Any, Dict, List, Optional, Tuple

from docx import Document

from .budget_sources import BudgetSource, normalize_budget_source
from .money import parse_money_to_rub


@dataclass
class ParsedSubprogram:
    number: int
    name: str
    status: str = "OK"


@dataclass
class ParsedProgramNode:
    stable_key: str
    node_type: str
    name: str
    parent_stable_key: Optional[str] = None
    display_number: Optional[str] = None
    code: Optional[str] = None
    normalized_name: Optional[str] = None
    execution_period: Optional[str] = None
    source_table_index: Optional[int] = None
    source_row_index: Optional[int] = None
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class ParsedFundingLine:
    node_stable_key: str
    year: int
    source_type: BudgetSource
    amount_rub: Decimal
    unit_in_document: str
    source_table_index: int
    source_row_index: int
    source_cell_index: int
    raw_value: str
    source_label: Optional[str] = None
    total_cell_index: Optional[int] = None
    total_raw_value: Optional[str] = None
    year_cell_indexes: Dict[str, int] = field(default_factory=dict)


@dataclass
class ParsedDocxProgram:
    program: Dict[str, Any] = field(default_factory=dict)
    subprograms: List[ParsedSubprogram] = field(default_factory=list)
    passport_amounts: Dict[Tuple[int, BudgetSource], Decimal] = field(default_factory=dict)
    passport_totals_by_year: Dict[int, Decimal] = field(default_factory=dict)
    passport_source_cell_coordinates: Dict[Tuple[int, BudgetSource], Dict[str, Any]] = field(default_factory=dict)
    passport_total_cell_coordinates: Dict[int, Dict[str, Any]] = field(default_factory=dict)
    passport_source_total_column_amounts: Dict[BudgetSource, Decimal] = field(default_factory=dict)
    passport_source_total_cell_coordinates: Dict[BudgetSource, Dict[str, Any]] = field(default_factory=dict)
    passport_grand_total_column_amount: Optional[Decimal] = None
    passport_grand_total_cell_coordinate: Optional[Dict[str, Any]] = None
    nodes: List[ParsedProgramNode] = field(default_factory=list)
    funding_lines: List[ParsedFundingLine] = field(default_factory=list)


def parse_docx_program(path: str | Path) -> ParsedDocxProgram:
    document = Document(path)
    parsed = ParsedDocxProgram()
    parsed.program = _extract_program_metadata(document)
    seen_nodes: set[str] = set()
    _append_node(
        parsed,
        seen_nodes,
        ParsedProgramNode(
            stable_key="program",
            node_type="program",
            name=parsed.program.get("name") or "Название не определено",
            normalized_name=_normalize_name(parsed.program.get("name") or "Название не определено"),
            metadata={"source": "docx_paragraphs"},
        ),
    )
    finance_table_index = 0
    for table_index, table in enumerate(document.tables):
        matrix = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        _extract_subprograms(matrix, parsed, table_index, seen_nodes)
        _extract_passport_amounts(matrix, parsed, table_index)
        _extract_result_nodes(matrix, parsed, table_index, seen_nodes)
        if _looks_like_finance_table(matrix):
            finance_table_index += 1
            _extract_finance_tree(matrix, parsed, table_index, finance_table_index, seen_nodes)
    return parsed


def _extract_program_metadata(document) -> Dict[str, Any]:
    metadata: Dict[str, Any] = {}
    for paragraph in document.paragraphs:
        text = _clean_cell(paragraph.text)
        lowered = text.lower()
        if "муницип" not in lowered or "программ" not in lowered:
            continue

        quoted = re.search(r"[«\"]([^»\"]+)[»\"]", text)
        if quoted and "name" not in metadata:
            metadata["name"] = quoted.group(1).strip()

        period = re.search(r"\b(20\d{2})\s*[-–]\s*(20\d{2})\b", text)
        if period:
            metadata["period_start_year"] = int(period.group(1))
            metadata["period_end_year"] = int(period.group(2))

        if metadata.get("name") and metadata.get("period_start_year"):
            break
    return metadata


def _extract_subprograms(matrix: List[List[str]], parsed: ParsedDocxProgram, table_index: int, seen_nodes: set[str]) -> None:
    for row_index, row in enumerate(matrix):
        joined = "\n".join(row)
        if "Подпрограмма" not in joined:
            continue
        for match in re.finditer(r"Подпрограмма\s+(\d+)\s+[«\"]?([^»\".\n]+)", joined):
            number = int(match.group(1))
            name = match.group(2).strip(" «»\".\n\t")
            if not any(existing.number == number for existing in parsed.subprograms):
                parsed.subprograms.append(ParsedSubprogram(number=number, name=name))
            _append_node(
                parsed,
                seen_nodes,
                ParsedProgramNode(
                    stable_key=_subprogram_key(number),
                    node_type="subprogram",
                    parent_stable_key="program",
                    display_number=str(number),
                    name=name,
                    normalized_name=_normalize_name(name),
                    source_table_index=table_index,
                    source_row_index=row_index,
                    metadata={"source": "subprogram_list"},
                ),
            )


def _extract_passport_amounts(matrix: List[List[str]], parsed: ParsedDocxProgram, table_index: int) -> None:
    if not matrix:
        return
    sample = " ".join(" ".join(row[:8]).lower() for row in matrix[:6])
    if "мероприят" in sample and "финанс" in sample:
        return
    header_index = None
    years_by_col: Dict[int, int] = {}
    total_col: Optional[int] = None
    for idx, row in enumerate(matrix):
        for col, cell in enumerate(row):
            match = re.search(r"\b(20\d{2})\b", cell.strip())
            if match:
                years_by_col[col] = int(match.group(1))
        if years_by_col and any("источник" in cell.lower() for cell in row):
            header_index = idx
            total_col = _passport_total_column(row, years_by_col)
            break
    if header_index is None:
        return
    for row_index, row in enumerate(matrix[header_index + 1 :], start=header_index + 1):
        if not row:
            continue
        if row[0].lower().strip().startswith("всего"):
            if total_col is not None and total_col < len(row):
                raw_total = _clean_cell(row[total_col])
                parsed.passport_grand_total_column_amount = parse_money_to_rub(raw_total, unit="thousand_rub")
                parsed.passport_grand_total_cell_coordinate = _cell_coordinate(
                    table_index,
                    row_index,
                    total_col,
                    raw_total,
                    source_cell_index=0,
                )
            for col, year in years_by_col.items():
                if col < len(row):
                    raw_value = _clean_cell(row[col])
                    parsed.passport_totals_by_year[year] = parse_money_to_rub(raw_value, unit="thousand_rub")
                    parsed.passport_total_cell_coordinates[year] = _cell_coordinate(
                        table_index,
                        row_index,
                        col,
                        raw_value,
                        source_cell_index=0,
                    )
            continue
        source = normalize_budget_source(row[0])
        if source is BudgetSource.UNKNOWN:
            continue
        if total_col is not None and total_col < len(row):
            raw_total = _clean_cell(row[total_col])
            parsed.passport_source_total_column_amounts[source] = parse_money_to_rub(raw_total, unit="thousand_rub")
            parsed.passport_source_total_cell_coordinates[source] = _cell_coordinate(
                table_index,
                row_index,
                total_col,
                raw_total,
                source_cell_index=0,
            )
        for col, year in years_by_col.items():
            if col < len(row):
                raw_value = _clean_cell(row[col])
                parsed.passport_amounts[(year, source)] = parse_money_to_rub(raw_value, unit="thousand_rub")
                parsed.passport_source_cell_coordinates[(year, source)] = _cell_coordinate(
                    table_index,
                    row_index,
                    col,
                    raw_value,
                    source_cell_index=0,
                )


def _passport_total_column(row: List[str], years_by_col: Dict[int, int]) -> Optional[int]:
    for col, cell in enumerate(row):
        if col in years_by_col:
            continue
        normalized = _normalize_name(cell)
        if normalized == "всего" or normalized.startswith("всего "):
            return col
    return None


def _extract_result_nodes(matrix: List[List[str]], parsed: ParsedDocxProgram, table_index: int, seen_nodes: set[str]) -> None:
    header_index = None
    for index, row in enumerate(matrix[:5]):
        lowered = [_clean_cell(cell).lower() for cell in row]
        if any("наименование результата" in cell for cell in lowered):
            header_index = index
            break
    if header_index is None:
        return

    for row_index, row in enumerate(matrix[header_index + 1 :], start=header_index + 1):
        if len(row) < 5:
            continue
        name = _clean_cell(row[4])
        if not name or name.lower().startswith("наименование"):
            continue
        subprogram_number = _safe_int(row[1])
        display_number = ".".join(part for part in [_clean_cell(row[2]), _clean_cell(row[3])] if part)
        parent_key = _subprogram_key(subprogram_number) if subprogram_number else "program"
        stable_key = f"result:{table_index}:{row_index}:{_slug(display_number or name)}"
        _append_node(
            parsed,
            seen_nodes,
            ParsedProgramNode(
                stable_key=stable_key,
                node_type="result",
                parent_stable_key=parent_key,
                display_number=display_number,
                code=_clean_cell(row[3]) or None,
                name=name,
                normalized_name=_normalize_name(name),
                source_table_index=table_index,
                source_row_index=row_index,
                metadata={
                    "source": "result_table",
                    "unit": _clean_cell(row[5]) if len(row) > 5 else "",
                    "calculation": _clean_cell(row[6]) if len(row) > 6 else "",
                },
            ),
        )


def _looks_like_finance_table(matrix: List[List[str]]) -> bool:
    if not matrix:
        return False
    sample = " ".join(" ".join(row[:6]).lower() for row in matrix[:4])
    return "источник" in sample and "финанс" in sample and ("мероприят" in sample or "подпрограмм" in sample)


def _extract_finance_tree(
    matrix: List[List[str]],
    parsed: ParsedDocxProgram,
    table_index: int,
    finance_table_index: int,
    seen_nodes: set[str],
) -> None:
    years_by_col = _finance_year_columns(matrix)
    if not years_by_col:
        return
    total_cell_index = _finance_total_column(matrix, years_by_col)

    data_start = _finance_data_start(matrix)
    subprogram_key = _subprogram_key_for_finance_table(parsed, finance_table_index)
    current_main_key: Optional[str] = None
    current_activity_key: Optional[str] = None
    node_by_signature: Dict[Tuple[str, str], str] = {}

    for row_index, row in enumerate(matrix[data_start:], start=data_start):
        if len(row) < 4:
            continue
        display_number = _normalize_display_number(row[0])
        name = _clean_cell(row[1])
        if not display_number:
            display_number = _infer_finance_activity_display_number(name)
        execution_period = _clean_cell(row[2])
        source_cell_index, row_total_cell_index, row_years_by_col = _finance_row_coordinates(
            row,
            years_by_col,
            total_cell_index,
        )
        year_cell_indexes = {str(year): col for col, year in row_years_by_col.items()}
        source_text = _clean_cell(row[source_cell_index]) if source_cell_index < len(row) else ""
        if not display_number or not name:
            continue

        source = normalize_budget_source(source_text)
        is_total_row = _is_total_source(source_text)
        signature = (display_number, _normalize_name(name))
        row_coordinate_metadata = _finance_row_coordinate_metadata(
            row,
            source_text,
            row_years_by_col,
            row_total_cell_index,
            source_cell_index,
            node_type=_classify_finance_node(display_number, name),
        )

        if is_total_row:
            node_type = _classify_finance_node(display_number, name)
            is_summary_row = _is_summary_total_node(display_number, name)
            if node_type == "main_activity":
                parent_key = subprogram_key
            elif node_type == "activity":
                parent_key = current_main_key or subprogram_key
            elif is_summary_row and _normalize_name(f"{display_number} {name}").find("подпрограмм") >= 0:
                parent_key = subprogram_key
            else:
                parent_key = current_activity_key or current_main_key or subprogram_key

            stable_key = f"{node_type}:{table_index}:{row_index}:{_slug(display_number)}"
            _append_node(
                parsed,
                seen_nodes,
                ParsedProgramNode(
                    stable_key=stable_key,
                    node_type=node_type,
                    parent_stable_key=parent_key,
                    display_number=display_number,
                    code=_extract_code(display_number, name),
                    name=name,
                    normalized_name=_normalize_name(name),
                    execution_period=execution_period or None,
                    source_table_index=table_index,
                    source_row_index=row_index,
                    metadata={
                        "source": "finance_table",
                        "finance_table_index": finance_table_index,
                        "docx_summary_row": is_summary_row,
                        **row_coordinate_metadata,
                    },
                ),
            )
            node_by_signature[signature] = stable_key
            if node_type == "main_activity":
                current_main_key = stable_key
                current_activity_key = None
            elif node_type == "activity":
                current_activity_key = stable_key
            continue

        if source is BudgetSource.UNKNOWN:
            continue

        node_key = node_by_signature.get(signature)
        if node_key is None:
            node_key = f"object:{table_index}:{row_index}:{_slug(display_number)}"
            _append_node(
                parsed,
                seen_nodes,
                ParsedProgramNode(
                    stable_key=node_key,
                    node_type="object",
                    parent_stable_key=current_activity_key or current_main_key or subprogram_key,
                    display_number=display_number,
                    code=_extract_code(display_number, name),
                    name=name,
                    normalized_name=_normalize_name(name),
                    execution_period=execution_period or None,
                    source_table_index=table_index,
                    source_row_index=row_index,
                    metadata={
                        "source": "finance_source_row",
                        "finance_table_index": finance_table_index,
                        "docx_row_type": "object",
                    },
                ),
            )
            node_by_signature[signature] = node_key

        for col, year in row_years_by_col.items():
            if col >= len(row):
                continue
            raw_value = _clean_cell(row[col])
            if not raw_value:
                continue
            parsed.funding_lines.append(
                ParsedFundingLine(
                    node_stable_key=node_key,
                    year=year,
                    source_type=source,
                    amount_rub=parse_money_to_rub(raw_value, unit="thousand_rub"),
                    unit_in_document="thousand_rub",
                    source_table_index=table_index,
                    source_row_index=row_index,
                    source_cell_index=col,
                    raw_value=raw_value,
                    source_label=source_text,
                    total_cell_index=row_total_cell_index,
                    total_raw_value=_clean_cell(row[row_total_cell_index]) if row_total_cell_index is not None and row_total_cell_index < len(row) else None,
                    year_cell_indexes=year_cell_indexes,
                )
            )


def _append_node(parsed: ParsedDocxProgram, seen_nodes: set[str], node: ParsedProgramNode) -> None:
    if node.stable_key in seen_nodes:
        return
    seen_nodes.add(node.stable_key)
    parsed.nodes.append(node)


def _cell_coordinate(
    table_index: int,
    row_index: int,
    cell_index: int,
    raw_value: str,
    source_cell_index: Optional[int] = None,
    unit: str = "thousand_rub",
) -> Dict[str, Any]:
    coordinate: Dict[str, Any] = {
        "table_index": table_index,
        "row_index": row_index,
        "cell_index": cell_index,
        "raw_value": raw_value,
        "unit_in_document": unit,
        "coordinate_key": f"{table_index}:{row_index}:{cell_index}",
    }
    if source_cell_index is not None:
        coordinate["source_cell_index"] = source_cell_index
    return coordinate


def _finance_row_coordinate_metadata(
    row: List[str],
    source_text: str,
    years_by_col: Dict[int, int],
    total_cell_index: Optional[int],
    source_cell_index: int,
    node_type: str,
) -> Dict[str, Any]:
    metadata: Dict[str, Any] = {
        "docx_row_type": node_type,
        "docx_source_cell_index": source_cell_index,
        "docx_source_raw_value": _clean_cell(source_text),
        "docx_year_cell_indexes": {str(year): col for col, year in years_by_col.items()},
        "docx_year_raw_values": {
            str(year): _clean_cell(row[col])
            for col, year in years_by_col.items()
            if col < len(row)
        },
        "docx_unit_in_document": "thousand_rub",
    }
    if total_cell_index is not None:
        metadata["docx_total_cell_index"] = total_cell_index
        metadata["docx_total_raw_value"] = _clean_cell(row[total_cell_index]) if total_cell_index < len(row) else ""
    return metadata


def _finance_row_coordinates(
    row: List[str],
    years_by_col: Dict[int, int],
    total_cell_index: Optional[int],
) -> Tuple[int, Optional[int], Dict[int, int]]:
    source_cell_index = 3
    row_total_cell_index = total_cell_index
    row_years_by_col = dict(years_by_col)
    if not _summary_row_uses_shifted_source(row):
        return source_cell_index, row_total_cell_index, row_years_by_col

    source_cell_index = 4
    if total_cell_index is not None:
        row_total_cell_index = total_cell_index + 1
        if row_total_cell_index in row_years_by_col:
            first_year = row_years_by_col[row_total_cell_index]
            shifted_year_cell_index = row_total_cell_index + 1
            if shifted_year_cell_index < len(row) and shifted_year_cell_index not in row_years_by_col:
                row_years_by_col.pop(row_total_cell_index)
                row_years_by_col[shifted_year_cell_index] = first_year
    return source_cell_index, row_total_cell_index, dict(sorted(row_years_by_col.items()))


def _summary_row_uses_shifted_source(row: List[str]) -> bool:
    if len(row) <= 4:
        return False
    display_number = _normalize_display_number(row[0])
    name = _clean_cell(row[1])
    if not _is_summary_total_node(display_number, name):
        return False
    first_cells = [_normalize_name(row[index]) for index in range(min(4, len(row)))]
    if not all(first_cells) or len(set(first_cells)) != 1:
        return False

    source_candidate = _clean_cell(row[4])
    if _is_total_source(source_candidate):
        return True
    if normalize_budget_source(source_candidate) is not BudgetSource.UNKNOWN:
        return True
    return "источник" in _normalize_name(source_candidate)


def _finance_year_columns(matrix: List[List[str]]) -> Dict[int, int]:
    result: Dict[int, int] = {}
    seen_years: set[int] = set()
    for row in matrix[:4]:
        for col, cell in enumerate(row):
            if col < 5:
                continue
            match = re.search(r"\b(20\d{2})\b", _clean_cell(cell))
            if not match:
                continue
            year = int(match.group(1))
            if year in seen_years:
                continue
            seen_years.add(year)
            result[col] = year
    return result


def _finance_total_column(matrix: List[List[str]], years_by_col: Dict[int, int]) -> Optional[int]:
    year_columns = set(years_by_col.keys())
    first_year_col = min(year_columns) if year_columns else None
    for row in matrix[:4]:
        for col, cell in enumerate(row):
            if col in year_columns:
                continue
            if first_year_col is not None and col > first_year_col:
                continue
            normalized = _normalize_name(cell)
            if normalized == "всего" or normalized.startswith("всего "):
                return col
    return None


def _finance_data_start(matrix: List[List[str]]) -> int:
    for index, row in enumerate(matrix[:6]):
        if len(row) > 3 and _clean_cell(row[0]) == "1" and _clean_cell(row[1]) == "2" and _clean_cell(row[3]) == "4":
            return index + 1
    for index, row in enumerate(matrix[:5]):
        if any("источник" in _clean_cell(cell).lower() for cell in row):
            return index + 1
    return 1


def _subprogram_key_for_finance_table(parsed: ParsedDocxProgram, finance_table_index: int) -> str:
    numbers = sorted(item.number for item in parsed.subprograms)
    if 0 < finance_table_index <= len(numbers):
        return _subprogram_key(numbers[finance_table_index - 1])
    return "program"


def _classify_finance_node(display_number: str, name: str) -> str:
    lowered = _normalize_name(name)
    if "основное мероприятие" in lowered:
        return "main_activity"
    if lowered.startswith("мероприятие") or re.match(r"^\d+(?:\.\d+)?$", display_number):
        return "activity"
    return "object"


def _infer_finance_activity_display_number(name: str) -> str:
    lowered = _normalize_name(name)
    if not lowered.startswith("мероприятие"):
        return ""
    match = re.search(r"\b(\d{2}(?:\.\d{2})?)\b", _clean_cell(name))
    return match.group(1) if match else ""


def _is_summary_total_node(display_number: str, name: str) -> bool:
    normalized = _normalize_name(f"{display_number} {name}")
    return bool(re.match(r"^(итого|всего)\b", normalized) or re.search(r"\b(итого|всего)\s+по\b", normalized))


def _extract_code(display_number: str, name: str) -> Optional[str]:
    match = re.search(r"\b(\d{2}(?:\.\d{2})?)\b", name)
    if match:
        return match.group(1)
    return display_number or None


def _is_total_source(text: str) -> bool:
    normalized = _normalize_name(text)
    return normalized in {"итого", "всего"} or normalized.startswith("итого ")


def _subprogram_key(number: Optional[int]) -> str:
    return f"subprogram:{number}" if number else "program"


def _safe_int(value: object) -> Optional[int]:
    match = re.search(r"\d+", _clean_cell(value))
    return int(match.group(0)) if match else None


def _normalize_display_number(value: object) -> str:
    text = _clean_cell(value)
    text = re.sub(r"\s*\|\s*", ".", text)
    text = re.sub(r"\s+", " ", text)
    text = re.sub(r"\.+", ".", text)
    return text.strip(" .")


def _clean_cell(value: object) -> str:
    text = "" if value is None else str(value)
    text = text.replace("\u00a0", " ").replace("\u202f", " ").replace("\u200b", "")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def _normalize_name(value: object) -> str:
    text = _clean_cell(value).lower().replace("ё", "е")
    text = re.sub(r"[«»\"“”]", "", text)
    return text.strip()


def _slug(value: object) -> str:
    text = _normalize_name(value)
    text = re.sub(r"[^0-9a-zа-я]+", "-", text)
    return text.strip("-") or "node"
