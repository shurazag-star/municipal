from decimal import Decimal
from typing import Optional

from docx import Document
from docx.oxml.ns import qn

from municipal_agent.docx_patcher import format_money_for_docx, patch_docx


def test_format_money_for_docx_preserves_precision_and_grouping():
    formatted = format_money_for_docx(
        Decimal("29163160.00"),
        source_cell_raw_value="1 000,00",
        unit="thousand_rub",
    )

    assert formatted == "29 163,16"


def test_format_money_for_docx_keeps_large_ungrouped_source_style_by_default():
    formatted = format_money_for_docx(
        Decimal("29163160.00"),
        source_cell_raw_value="1000,00",
        unit="thousand_rub",
    )

    assert formatted == "29163,16"


def test_format_money_for_docx_keeps_small_ungrouped_source_style():
    formatted = format_money_for_docx(
        Decimal("963900.00"),
        source_cell_raw_value="1000,00",
        unit="thousand_rub",
    )

    assert formatted == "963,90"


def test_patch_docx_replaces_old_approval_date_with_placeholders(tmp_path):
    source_path = tmp_path / "source_approval.docx"
    output_path = tmp_path / "generated_approval.docx"
    document = Document()
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "УТВЕРЖДЕНА Постановлением администрации от 20.02.2026 № 658-ПА"
    document.save(source_path)

    result = patch_docx(
        input_path=source_path,
        output_path=output_path,
        changes={"cell_updates": [], "text_updates": [], "insert_objects": []},
    )

    generated = Document(output_path)
    assert result["approval_header_normalized"] is True
    assert generated.tables[0].cell(0, 0).text == "УТВЕРЖДЕНА Постановлением администрации от _______________ №__________"


def test_patch_docx_updates_numeric_cell_and_preserves_source_file(tmp_path):
    source_path = tmp_path / "source.docx"
    output_path = tmp_path / "generated.docx"
    document = Document()
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Источник"
    table.cell(0, 1).text = "2026"
    table.cell(1, 0).text = "Местный бюджет"
    table.cell(1, 1).text = "100,00"
    document.save(source_path)
    original_bytes = source_path.read_bytes()

    result = patch_docx(
        input_path=source_path,
        output_path=output_path,
        changes=[
            {
                "change_item_id": 1,
                "table_index": 0,
                "row_index": 1,
                "cell_index": 1,
                "amount_rub": "150000.00",
                "source_cell_raw_value": "100,00",
                "unit": "thousand_rub",
            }
        ],
    )

    assert result["applied_count"] == 1
    assert result["skipped_count"] == 0
    assert source_path.read_bytes() == original_bytes
    generated = Document(output_path)
    assert generated.tables[0].cell(1, 1).text == "150,00"


def test_patch_docx_uses_visual_row_cells_for_merged_tables(tmp_path, monkeypatch):
    source_path = tmp_path / "source.docx"
    output_path = tmp_path / "generated.docx"
    source_path.write_bytes(b"fake docx")

    correct_cell = FakeCell("0,0")
    wrong_cell = FakeCell("wrong")
    document = FakeDocument(
        FakeTable(
            rows=[
                FakeRow([FakeCell("header")]),
                FakeRow([FakeCell("source"), correct_cell]),
            ],
            wrong_cell=wrong_cell,
        )
    )
    monkeypatch.setattr("municipal_agent.docx_patcher.Document", lambda _path: document)

    result = patch_docx(
        input_path=source_path,
        output_path=output_path,
        changes=[
            {
                "change_item_id": 2,
                "table_index": 0,
                "row_index": 1,
                "cell_index": 1,
                "amount_rub": "17898180.00",
                "source_cell_raw_value": "0,0",
                "unit": "thousand_rub",
            }
        ],
    )

    assert result["applied_count"] == 1
    assert correct_cell.text == "17898,18"
    assert wrong_cell.text == "wrong"
    assert document.saved_path == str(output_path)


def test_patch_docx_inserts_new_object_rows_from_template(tmp_path):
    source_path = tmp_path / "source.docx"
    output_path = tmp_path / "generated.docx"
    document = Document()
    table = document.add_table(rows=3, cols=7)
    headers = ["N", "Наименование", "Период", "Источник", "Всего", "2026", "2027"]
    for index, header in enumerate(headers):
        table.cell(0, index).text = header
    table.cell(1, 0).text = "2.1.1"
    table.cell(1, 1).text = "Существующий объект"
    table.cell(1, 2).text = "2026-2030"
    table.cell(1, 3).text = "Итого"
    table.cell(1, 4).text = "100,00"
    table.cell(1, 5).text = "40,00"
    table.cell(1, 6).text = "60,00"
    table.cell(2, 0).text = "2.1.1."
    table.cell(2, 1).text = "Существующий объект"
    table.cell(2, 2).text = "2026-2030"
    table.cell(2, 3).text = "Местный бюджет"
    table.cell(2, 4).text = "100,00"
    table.cell(2, 5).text = "40,00"
    table.cell(2, 6).text = "60,00"
    document.save(source_path)
    original_bytes = source_path.read_bytes()

    result = patch_docx(
        input_path=source_path,
        output_path=output_path,
        changes={
            "cell_updates": [],
            "insert_objects": [
                {
                    "change_item_ids": [10, 11],
                    "table_index": 0,
                    "insert_after_row_index": 2,
                    "template_row_index": 2,
                    "display_number": "2.1.2",
                    "object_name": "Новый объект водоснабжения",
                    "execution_period": "2026-2027",
                    "total_cell_index": 4,
                    "year_cell_indices": {"2026": 5, "2027": 6},
                    "rows": [
                        {
                            "source_type": "TOTAL",
                            "source_label": "Итого",
                            "total_amount_rub": "300000.00",
                            "amounts_by_year": {"2026": "100000.00", "2027": "200000.00"},
                            "unit": "thousand_rub",
                        },
                        {
                            "source_type": "LOCAL_BUDGET",
                            "source_label": "Местный бюджет",
                            "total_amount_rub": "300000.00",
                            "amounts_by_year": {"2026": "100000.00", "2027": "200000.00"},
                            "unit": "thousand_rub",
                        },
                    ],
                }
            ],
        },
    )

    assert result["inserted_count"] == 1
    assert result["skipped_insertions_count"] == 0
    assert source_path.read_bytes() == original_bytes

    generated = Document(output_path)
    table = generated.tables[0]
    assert len(table.rows) == 5
    assert table.cell(3, 0).text == "2.1.2."
    assert table.cell(3, 1).text == "Новый объект водоснабжения"
    assert table.cell(3, 3).text == "Итого"
    assert table.cell(3, 4).text == "300,00"
    assert table.cell(3, 5).text == "100,00"
    assert table.cell(3, 6).text == "200,00"
    assert table.cell(4, 3).text == "Местный бюджет"


def test_format_money_groups_inserted_values_without_template_grouping():
    assert (
        format_money_for_docx("11245920.00", source_cell_raw_value="0,00", unit="thousand_rub", default_grouping=True)
        == "11 245,92"
    )


def test_patch_docx_infers_inserted_row_responsible_from_next_row(tmp_path):
    source_path = tmp_path / "source_responsible.docx"
    output_path = tmp_path / "generated_responsible.docx"
    document = Document()
    table = document.add_table(rows=4, cols=8)
    headers = ["N", "Наименование", "Период", "Источник", "Всего", "2026", "2027", "Исполнитель"]
    for index, header in enumerate(headers):
        table.cell(0, index).text = header

    table.cell(1, 0).text = "1.1."
    table.cell(1, 1).text = "Предыдущее мероприятие"
    table.cell(1, 2).text = "2026-2027"
    table.cell(1, 3).text = "Итого:"
    table.cell(1, 7).text = "Предыдущий исполнитель"
    table.cell(2, 0).text = "1.1."
    table.cell(2, 1).text = "Предыдущее мероприятие"
    table.cell(2, 2).text = "2026-2027"
    table.cell(2, 3).text = "Местный бюджет"
    table.cell(2, 7).text = "Предыдущий исполнитель"
    table.cell(3, 0).text = "1.3."
    table.cell(3, 1).text = "Следующее мероприятие"
    table.cell(3, 2).text = "2026-2027"
    table.cell(3, 3).text = "Местный бюджет"
    table.cell(3, 7).text = "Управление молодежной политики"
    document.save(source_path)

    result = patch_docx(
        input_path=source_path,
        output_path=output_path,
        changes={
            "cell_updates": [],
            "insert_objects": [
                {
                    "table_index": 0,
                    "insert_after_row_index": 2,
                    "template_row_index": 2,
                    "display_number": "1.2",
                    "object_name": "Мероприятие 01.02. Новое мероприятие",
                    "execution_period": "2026-2027",
                    "total_cell_index": 4,
                    "year_cell_indices": {"2026": 5, "2027": 6},
                    "rows": [
                        {
                            "source_type": "TOTAL",
                            "source_label": "Итого:",
                            "total_amount_rub": "150000.00",
                            "amounts_by_year": {"2026": "150000.00", "2027": "0.00"},
                            "unit": "thousand_rub",
                        },
                        {
                            "source_type": "LOCAL_BUDGET",
                            "source_label": "Местный бюджет",
                            "total_amount_rub": "150000.00",
                            "amounts_by_year": {"2026": "150000.00", "2027": "0.00"},
                            "unit": "thousand_rub",
                        }
                    ],
                }
            ],
        },
    )

    assert result["inserted_count"] == 1
    generated = Document(output_path)
    table = generated.tables[0]
    assert table.cell(3, 0).text == "1.2."
    assert table.cell(3, 7).text == "Управление молодежной политики"
    assert table.cell(4, 7).text == "Управление молодежной политики"


def test_patch_docx_applies_text_updates(tmp_path):
    source_path = tmp_path / "source_text.docx"
    output_path = tmp_path / "generated_text.docx"
    document = Document()
    table = document.add_table(rows=2, cols=3)
    table.cell(1, 0).text = "2.1.1"
    table.cell(1, 1).text = "Объект"
    table.cell(1, 2).text = "2026-2028"
    document.save(source_path)

    result = patch_docx(
        input_path=source_path,
        output_path=output_path,
        changes={
            "cell_updates": [],
            "text_updates": [
                {
                    "table_index": 0,
                    "row_index": 1,
                    "cell_index": 2,
                    "text": "2027-2028",
                }
            ],
            "insert_objects": [],
        },
    )

    assert result["text_applied_count"] == 1
    generated = Document(output_path)
    assert generated.tables[0].cell(1, 2).text == "2027-2028"


def test_patch_docx_increments_parent_result_count_for_inserted_object(tmp_path):
    source_path = tmp_path / "source_result.docx"
    output_path = tmp_path / "generated_result.docx"
    document = Document()
    table = document.add_table(rows=5, cols=8)
    table.rows[0].cells[0].text = "2.1"
    table.rows[0].cells[1].text = "Мероприятие"
    table.rows[1].cells[0].text = "2.1"
    table.rows[1].cells[1].text = "Построены объекты"
    table.rows[1].cells[2].text = "Х"
    table.rows[1].cells[3].text = "Х"
    table.rows[1].cells[4].text = "Всего"
    table.rows[1].cells[5].text = "2026 год"
    table.rows[1].cells[6].text = "2027"
    table.rows[2].cells[0].text = "2.1"
    table.rows[2].cells[1].text = "Построены объекты"
    table.rows[2].cells[2].text = "Х"
    table.rows[2].cells[3].text = "Х"
    table.rows[2].cells[4].text = "7"
    table.rows[2].cells[5].text = "1"
    table.rows[2].cells[6].text = "6"
    table.rows[3].cells[0].text = "2.1.1"
    table.rows[3].cells[1].text = "Существующий объект"
    table.rows[3].cells[2].text = "2027"
    table.rows[3].cells[3].text = "Итого"
    table.rows[3].cells[4].text = "100,00"
    table.rows[3].cells[6].text = "100,00"
    document.save(source_path)

    result = patch_docx(
        input_path=source_path,
        output_path=output_path,
        changes={
            "cell_updates": [],
            "insert_objects": [
                {
                    "table_index": 0,
                    "insert_after_row_index": 3,
                    "template_row_index": 3,
                    "parent_display_number": "2.1",
                    "display_number": "2.1.2",
                    "object_name": "Новый объект",
                    "execution_period": "2027",
                    "active_years": [2027],
                    "total_cell_index": 4,
                    "year_cell_indices": {"2027": 6},
                    "rows": [
                        {
                            "source_type": "TOTAL",
                            "source_label": "Итого",
                            "total_amount_rub": "200000.00",
                            "amounts_by_year": {"2027": "200000.00"},
                        }
                    ],
                }
            ],
        },
    )

    assert result["result_count_applied_count"] == 1
    generated = Document(output_path)
    row = generated.tables[0].rows[2].cells
    assert row[4].text == "8"
    assert row[5].text == "1"
    assert row[6].text == "7"


def test_patch_docx_inserts_after_vertically_merged_object_without_rebinding_old_rows(tmp_path):
    source_path = tmp_path / "source_merged.docx"
    output_path = tmp_path / "generated_merged.docx"
    document = Document()
    table = document.add_table(rows=4, cols=7)
    headers = ["N", "Наименование", "Период", "Источник", "Всего", "2026", "2027"]
    for index, header in enumerate(headers):
        table.cell(0, index).text = header

    table.cell(1, 0).text = "2.1.1"
    table.cell(1, 1).text = "Существующий объект"
    table.cell(1, 2).text = "2026-2027"
    for col in [0, 1, 2]:
        table.cell(1, col).merge(table.cell(3, col))
    table.cell(1, 3).text = "Итого"
    table.cell(2, 3).text = "Средства бюджета Московской области"
    table.cell(3, 3).text = "Средства бюджета муниципального округа Шатура"
    for row_index, value in [(1, "300,00"), (2, "200,00"), (3, "100,00")]:
        table.cell(row_index, 4).text = value
        table.cell(row_index, 5).text = value
        table.cell(row_index, 6).text = "0,00"
    document.save(source_path)

    result = patch_docx(
        input_path=source_path,
        output_path=output_path,
        changes={
            "cell_updates": [],
            "insert_objects": [
                {
                    "change_item_ids": [20],
                    "table_index": 0,
                    "insert_after_row_index": 3,
                    "template_row_index": 1,
                    "display_number": "2.1.2",
                    "object_name": "Новый объект",
                    "execution_period": "2026",
                    "total_cell_index": 4,
                    "year_cell_indices": {"2026": 5, "2027": 6},
                    "rows": [
                        {
                            "source_type": "TOTAL",
                            "source_label": "Итого",
                            "total_amount_rub": "150000.00",
                            "amounts_by_year": {"2026": "150000.00", "2027": "0.00"},
                            "unit": "thousand_rub",
                        },
                        {
                            "source_type": "LOCAL_BUDGET",
                            "source_label": "Средства бюджета Шатура",
                            "total_amount_rub": "150000.00",
                            "amounts_by_year": {"2026": "150000.00", "2027": "0.00"},
                            "unit": "thousand_rub",
                        },
                    ],
                }
            ],
        },
    )

    assert result["inserted_count"] == 1
    assert result["skipped_insertions_count"] == 0

    generated = Document(output_path)
    table = generated.tables[0]
    assert len(table.rows) == 6
    assert table.cell(1, 1).text == "Существующий объект"
    assert table.cell(3, 3).text == "Средства бюджета муниципального округа Шатура"
    assert table.cell(4, 1).text == "Новый объект"
    assert table.cell(4, 3).text == "Итого"
    assert table.cell(5, 1).text == "Новый объект"
    assert table.cell(5, 3).text == "Средства бюджета Шатура"
    assert _vmerge_value(table, 4, 0) == "restart"
    assert _vmerge_value(table, 5, 0) == "continue"
    assert _vmerge_value(table, 4, 1) == "restart"
    assert _vmerge_value(table, 5, 1) == "continue"
    assert _vmerge_value(table, 4, 2) == "restart"
    assert _vmerge_value(table, 5, 2) == "continue"
    assert _raw_cell_text(table, 5, 1) == ""


def _vmerge_value(table, row_index: int, cell_index: int) -> Optional[str]:
    tc = table._tbl.tr_lst[row_index].tc_lst[cell_index]
    properties = tc.tcPr
    if properties is None:
        return None

    vertical_merge = properties.find(qn("w:vMerge"))
    if vertical_merge is None:
        return None

    return vertical_merge.get(qn("w:val")) or "continue"


def _raw_cell_text(table, row_index: int, cell_index: int) -> str:
    tc = table._tbl.tr_lst[row_index].tc_lst[cell_index]
    return "".join(node.text or "" for node in tc.iter(qn("w:t")))


class FakeRun:
    def __init__(self, text=""):
        self.text = text


class FakeParagraph:
    def __init__(self, text=""):
        self.runs = [FakeRun(text)]

    def add_run(self, value):
        run = FakeRun(value)
        self.runs.append(run)
        return run


class FakeCell:
    def __init__(self, text=""):
        self.paragraphs = [FakeParagraph(text)]

    @property
    def text(self):
        return "".join(run.text for paragraph in self.paragraphs for run in paragraph.runs)

    def add_paragraph(self):
        paragraph = FakeParagraph("")
        self.paragraphs.append(paragraph)
        return paragraph


class FakeRow:
    def __init__(self, cells):
        self.cells = cells


class FakeTable:
    def __init__(self, rows, wrong_cell):
        self.rows = rows
        self._wrong_cell = wrong_cell

    def cell(self, _row_index, _cell_index):
        return self._wrong_cell


class FakeDocument:
    def __init__(self, table):
        self.tables = [table]
        self.saved_path = None

    def save(self, path):
        self.saved_path = str(path)
