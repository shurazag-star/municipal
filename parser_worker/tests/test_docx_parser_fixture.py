from decimal import Decimal

from docx import Document

from municipal_agent.budget_sources import BudgetSource
from municipal_agent.docx_parser import parse_docx_program


def test_docx_parser_extracts_passport_subprograms_and_amounts(tmp_path):
    path = tmp_path / "program.docx"
    document = Document()
    document.add_heading("Муниципальная программа", level=1)
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Перечень подпрограмм"
    table.rows[0].cells[1].text = "\n".join(
        [
            "Подпрограмма 1 «Чистая вода».",
            "Подпрограмма 2 «Системы водоотведения».",
            "Подпрограмма 3 «Объекты теплоснабжения, инженерные коммуникации».",
            "Подпрограмма 4 «Обращение с отходами».",
            "Подпрограмма 5 «Энергосбережение и повышение энергетической эффективности».",
            "Подпрограмма 6 «Развитие газификации, топливнозаправочного комплекса и электроэнергетики».",
            "Подпрограмма 7 «Обеспечивающая подпрограмма».",
            "Подпрограмма 8 «Реализация полномочий в сфере жилищно-коммунального хозяйства».",
        ]
    )
    amounts = document.add_table(rows=1, cols=7)
    for idx, label in enumerate(["Источник", "Всего", "2026", "2027", "2028", "2029", "2030"]):
        amounts.rows[0].cells[idx].text = label
    row = amounts.add_row().cells
    row[0].text = "Бюджет Московской области"
    row[1].text = "3 736 797,28"
    row[2].text = "1 858 904,68"
    row[3].text = "1 344 541,65"
    row[4].text = "533 350,95"
    row[5].text = "0,00"
    row[6].text = "0,00"
    document.save(path)

    parsed = parse_docx_program(path)

    assert len(parsed.subprograms) == 8
    assert parsed.subprograms[0].name == "Чистая вода"
    assert parsed.passport_amounts[(2026, BudgetSource.MOSCOW_OBLAST_BUDGET)] == Decimal("1858904680.00")
    assert parsed.passport_amounts[(2027, BudgetSource.MOSCOW_OBLAST_BUDGET)] == Decimal("1344541650.00")
    assert parsed.passport_source_cell_coordinates[(2026, BudgetSource.MOSCOW_OBLAST_BUDGET)] == {
        "table_index": 1,
        "row_index": 1,
        "source_cell_index": 0,
        "cell_index": 2,
        "raw_value": "1 858 904,68",
        "unit_in_document": "thousand_rub",
        "coordinate_key": "1:1:2",
    }
    assert parsed.passport_source_total_column_amounts[BudgetSource.MOSCOW_OBLAST_BUDGET] == Decimal("3736797280.00")
    assert parsed.passport_source_total_cell_coordinates[BudgetSource.MOSCOW_OBLAST_BUDGET] == {
        "table_index": 1,
        "row_index": 1,
        "source_cell_index": 0,
        "cell_index": 1,
        "raw_value": "3 736 797,28",
        "unit_in_document": "thousand_rub",
        "coordinate_key": "1:1:1",
    }


def test_docx_parser_ignores_non_financing_source_words_before_passport(tmp_path):
    path = tmp_path / "program_with_source_word.docx"
    document = Document()
    table = document.add_table(rows=1, cols=7)
    table.rows[0].cells[0].text = "Сроки реализации муниципальной программы"
    for idx in range(1, 7):
        table.rows[0].cells[idx].text = "2026-2030"

    description = table.add_row().cells
    description[0].text = "Краткая характеристика подпрограмм"
    for idx in range(1, 7):
        description[idx].text = "Увеличение доли населения из централизованных источников водоснабжения"

    header = table.add_row().cells
    header[0].text = "Источники финансирования муниципальной программы, в том числе по годам реализации программы (тыс. руб.):"
    header[1].text = "Всего"
    header[2].text = "2026"
    header[3].text = "2027"
    header[4].text = "2028"
    header[5].text = "2029"
    header[6].text = "2030"

    local = table.add_row().cells
    local[0].text = "Средства бюджета Городского округа Люберцы"
    local[1].text = "585 651,52"
    local[2].text = "377 165,32"
    local[3].text = "208 486,20"
    local[4].text = "0,00"
    local[5].text = "0,00"
    local[6].text = "0,00"

    total = table.add_row().cells
    total[0].text = "Всего, в том числе по годам:"
    total[1].text = "2 300 537,01"
    total[2].text = "1 627 276,41"
    total[3].text = "673 260,60"
    total[4].text = "0,00"
    total[5].text = "0,00"
    total[6].text = "0,00"
    document.save(path)

    parsed = parse_docx_program(path)

    assert parsed.passport_totals_by_year == {
        2026: Decimal("1627276410.00"),
        2027: Decimal("673260600.00"),
        2028: Decimal("0.00"),
        2029: Decimal("0.00"),
        2030: Decimal("0.00"),
    }
    assert parsed.passport_total_cell_coordinates[2027]["cell_index"] == 3
    assert parsed.passport_source_total_column_amounts[BudgetSource.LOCAL_BUDGET] == Decimal("585651520.00")
    assert parsed.passport_source_total_cell_coordinates[BudgetSource.LOCAL_BUDGET]["cell_index"] == 1


def test_docx_parser_extracts_tree_nodes_and_funding_lines_with_coordinates(tmp_path):
    path = tmp_path / "tree_program.docx"
    document = Document()
    document.add_paragraph("Муниципальная программа «Развитие инженерной инфраструктуры» на 2026-2030 годы")
    subprogram_table = document.add_table(rows=1, cols=2)
    subprogram_table.rows[0].cells[0].text = "Перечень подпрограмм"
    subprogram_table.rows[0].cells[1].text = "Подпрограмма 1 «Чистая вода»"

    results_table = document.add_table(rows=1, cols=7)
    for idx, label in enumerate(
        [
            "№пп",
            "№подпрограммы",
            "№ основного мероприятия",
            "№ мероприятия",
            "Наименование результата",
            "Единица измерения",
            "Порядок определения значений",
        ]
    ):
        results_table.rows[0].cells[idx].text = label
    result_cells = results_table.add_row().cells
    result_cells[0].text = "1"
    result_cells[1].text = "1"
    result_cells[2].text = "02"
    result_cells[3].text = "02.01"
    result_cells[4].text = "Построены объекты водоснабжения"
    result_cells[5].text = "единица"
    result_cells[6].text = "По данным ОМСУ"

    finance_table = document.add_table(rows=3, cols=8)
    for idx, label in enumerate(["№ п/п", "Мероприятие подпрограммы", "Срок исполнения мероприятия", "Источники финансирования", "Всего", "2026 год", "2027 год", "2028год"]):
        finance_table.rows[0].cells[idx].text = label
        finance_table.rows[1].cells[idx].text = label
    finance_table.rows[2].cells[0].text = "1"
    finance_table.rows[2].cells[1].text = "Основное мероприятие 02 - Строительство объектов водоснабжения"
    finance_table.rows[2].cells[2].text = "2026-2027"
    finance_table.rows[2].cells[3].text = "Итого:"
    finance_table.rows[2].cells[4].text = "90 555,38"
    finance_table.rows[2].cells[5].text = "90 555,38"
    finance_table.rows[2].cells[6].text = "0,00"
    finance_table.rows[2].cells[7].text = "0,00"

    activity = finance_table.add_row().cells
    activity[0].text = "1.1"
    activity[1].text = "Мероприятие 02.01 - Строительство водозаборных узлов"
    activity[2].text = "2026-2027"
    activity[3].text = "Итого:"
    activity[4].text = "90 555,38"
    activity[5].text = "90 555,38"
    activity[6].text = "0,00"
    activity[7].text = "0,00"

    object_total = finance_table.add_row().cells
    object_total[0].text = "1.1.1"
    object_total[1].text = "ВЗУ Черусти"
    object_total[2].text = "2026"
    object_total[3].text = "Итого:"
    object_total[4].text = "90 555,38"
    object_total[5].text = "90 555,38"
    object_total[6].text = "0,00"
    object_total[7].text = "0,00"

    oblast = finance_table.add_row().cells
    oblast[0].text = "1.1.1"
    oblast[1].text = "ВЗУ Черусти"
    oblast[2].text = "2026"
    oblast[3].text = "Средства бюджета Московской области"
    oblast[5].text = "78 330,39"
    oblast[7].text = "1,00"

    local = finance_table.add_row().cells
    local[0].text = "1.1.1"
    local[1].text = "ВЗУ Черусти"
    local[2].text = "2026"
    local[3].text = "Средства бюджета муниципального округа"
    local[5].text = "12 224,99"
    local[7].text = "2,00"

    subprogram_total = finance_table.add_row().cells
    subprogram_total[0].text = "Итого по подпрограмме"
    subprogram_total[1].text = "Итого по подпрограмме"
    subprogram_total[3].text = "Итого:"
    subprogram_total[4].text = "90 555,38"
    subprogram_total[5].text = "90 555,38"
    subprogram_total[6].text = "0,00"
    subprogram_total[7].text = "0,00"

    blank_number_activity = finance_table.add_row().cells
    blank_number_activity[1].text = "Мероприятие 02.02 - Ремонт станции водоподготовки"
    blank_number_activity[3].text = "Итого:"
    blank_number_activity[4].text = "15 000,00"
    blank_number_activity[5].text = "15 000,00"
    blank_number_activity[6].text = "0,00"

    blank_number_local = finance_table.add_row().cells
    blank_number_local[1].text = "Мероприятие 02.02 - Ремонт станции водоподготовки"
    blank_number_local[3].text = "Средства бюджета муниципального округа"
    blank_number_local[5].text = "15 000,00"
    document.save(path)

    parsed = parse_docx_program(path)

    assert parsed.program["name"] == "Развитие инженерной инфраструктуры"
    assert parsed.program["period_start_year"] == 2026
    assert parsed.program["period_end_year"] == 2030
    assert {"program", "subprogram", "main_activity", "activity", "object", "result"}.issubset({node.node_type for node in parsed.nodes})
    object_node = next(node for node in parsed.nodes if node.node_type == "object" and "Черусти" in node.name)
    assert object_node.source_table_index == 2
    assert object_node.source_row_index == 4
    assert object_node.parent_stable_key
    assert object_node.metadata["docx_row_type"] == "object"
    assert object_node.metadata["docx_source_cell_index"] == 3
    assert object_node.metadata["docx_total_cell_index"] == 4
    assert object_node.metadata["docx_year_cell_indexes"] == {"2026": 5, "2027": 6, "2028": 7}
    summary_node = next(node for node in parsed.nodes if node.name == "Итого по подпрограмме")
    assert summary_node.metadata["docx_summary_row"] is True
    assert summary_node.parent_stable_key == "subprogram:1"
    blank_number_node = next(node for node in parsed.nodes if node.code == "02.02")
    assert blank_number_node.node_type == "activity"
    assert blank_number_node.display_number == "02.02"
    assert blank_number_node.parent_stable_key

    funding = [line for line in parsed.funding_lines if line.node_stable_key == object_node.stable_key]
    assert len(funding) == 4
    oblast_line = next(line for line in funding if line.source_type == BudgetSource.MOSCOW_OBLAST_BUDGET)
    assert oblast_line.amount_rub == Decimal("78330390.00")
    assert oblast_line.unit_in_document == "thousand_rub"
    assert oblast_line.source_table_index == 2
    assert oblast_line.source_row_index == 5
    assert oblast_line.source_cell_index == 5
    assert oblast_line.total_cell_index == 4
    assert oblast_line.year_cell_indexes == {"2026": 5, "2027": 6, "2028": 7}
    assert oblast_line.raw_value == "78 330,39"
    assert oblast_line.source_label == "Средства бюджета Московской области"
    local_2028 = next(line for line in funding if line.year == 2028 and line.source_type == BudgetSource.LOCAL_BUDGET)
    assert local_2028.amount_rub == Decimal("2000.00")
    assert local_2028.source_cell_index == 7
    blank_number_source_node = next(
        node
        for node in parsed.nodes
        if node.code == "02.02" and node.node_type == "object"
    )
    assert blank_number_source_node.parent_stable_key
    assert blank_number_source_node.parent_stable_key != blank_number_node.stable_key
    blank_number_funding = [
        line
        for line in parsed.funding_lines
        if line.node_stable_key == blank_number_source_node.stable_key
    ]
    assert len(blank_number_funding) == 1
    assert blank_number_funding[0].source_type == BudgetSource.LOCAL_BUDGET
    assert blank_number_funding[0].amount_rub == Decimal("15000000.00")


def test_docx_parser_extracts_shifted_summary_source_rows(tmp_path):
    path = tmp_path / "summary_rows.docx"
    document = Document()
    document.add_paragraph("Муниципальная программа «Развитие инженерной инфраструктуры» на 2026-2030 годы")
    subprogram_table = document.add_table(rows=1, cols=2)
    subprogram_table.rows[0].cells[0].text = "Перечень подпрограмм"
    subprogram_table.rows[0].cells[1].text = "Подпрограмма 1 «Системы водоотведения»"

    finance_table = document.add_table(rows=3, cols=20)
    headers = [
        "№ п/п",
        "Мероприятие подпрограммы",
        "Срок исполнения мероприятия",
        "Источники финансирования",
        "Всего (тыс. руб.)",
        "Объем финансирования по годам (тыс.руб.)",
        "Объем финансирования по годам (тыс.руб.)",
        "Объем финансирования по годам (тыс.руб.)",
        "Объем финансирования по годам (тыс.руб.)",
        "Объем финансирования по годам (тыс.руб.)",
        "Объем финансирования по годам (тыс.руб.)",
        "Объем финансирования по годам (тыс.руб.)",
        "Объем финансирования по годам (тыс.руб.)",
        "Объем финансирования по годам (тыс.руб.)",
        "Объем финансирования по годам (тыс.руб.)",
        "Объем финансирования по годам (тыс.руб.)",
        "Объем финансирования по годам (тыс.руб.)",
        "Объем финансирования по годам (тыс.руб.)",
        "Объем финансирования по годам (тыс.руб.)",
        "Ответственный",
    ]
    years = [
        "№ п/п",
        "Мероприятие подпрограммы",
        "Срок исполнения мероприятия",
        "Источники финансирования",
        "Всего (тыс. руб.)",
        "2026 год",
        "2026 год",
        "2026 год",
        "2026 год",
        "2026 год",
        "2026 год",
        "2026 год",
        "2026 год",
        "2026 год",
        "2026 год",
        "2027 год",
        "2028 год",
        "2029 год",
        "2030 год",
        "",
    ]
    for idx, label in enumerate(headers):
        finance_table.rows[0].cells[idx].text = label
    for idx, label in enumerate(years):
        finance_table.rows[1].cells[idx].text = label
        finance_table.rows[2].cells[idx].text = str(idx + 1)

    summary_total = finance_table.add_row().cells
    for idx in range(4):
        summary_total[idx].text = "Итого по подпрограмме"
    summary_total[4].text = "Итого:"
    summary_total[5].text = "759 864,58"
    summary_total[6].text = "68 810,11"
    summary_total[15].text = "261 756,23"
    summary_total[16].text = "429 298,24"

    summary_regional = finance_table.add_row().cells
    for idx in range(4):
        summary_regional[idx].text = "Итого по подпрограмме"
    summary_regional[4].text = "Средства бюджета Московской области"
    summary_regional[5].text = "592 986,49"
    summary_regional[6].text = "49 406,92"
    summary_regional[15].text = "204 016,95"
    summary_regional[16].text = "339 562,62"

    summary_local = finance_table.add_row().cells
    for idx in range(4):
        summary_local[idx].text = "Итого по подпрограмме"
    summary_local[4].text = "Средства бюджета муниципального округа Шатура"
    summary_local[5].text = "166 878,09"
    summary_local[6].text = "19 403,19"
    summary_local[15].text = "57 739,28"
    summary_local[16].text = "89 735,62"
    document.save(path)

    parsed = parse_docx_program(path)

    summary_node = next(node for node in parsed.nodes if node.name == "Итого по подпрограмме")
    assert summary_node.metadata["docx_summary_row"] is True
    assert summary_node.metadata["docx_source_cell_index"] == 4
    assert summary_node.metadata["docx_total_cell_index"] == 5
    assert summary_node.metadata["docx_year_cell_indexes"] == {"2026": 6, "2027": 15, "2028": 16, "2029": 17, "2030": 18}
    funding = [line for line in parsed.funding_lines if line.node_stable_key == summary_node.stable_key]
    assert len(funding) == 6
    regional_2026 = next(line for line in funding if line.year == 2026 and line.source_type == BudgetSource.MOSCOW_OBLAST_BUDGET)
    local_2028 = next(line for line in funding if line.year == 2028 and line.source_type == BudgetSource.LOCAL_BUDGET)
    assert regional_2026.amount_rub == Decimal("49406920.00")
    assert regional_2026.total_cell_index == 5
    assert regional_2026.year_cell_indexes == {"2026": 6, "2027": 15, "2028": 16, "2029": 17, "2030": 18}
    assert local_2028.amount_rub == Decimal("89735620.00")
